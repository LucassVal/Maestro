#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import sys
from .extrator_base_SRC_GL import ExtratorBase

try:
    from docx import Document
except ImportError:
    Document = None

class ExtratorDOCX(ExtratorBase):
    def __init__(self):
        super().__init__(nome="ExtratorDOCX")

    def executar(self, entrada):
        caminho = entrada.get("caminho")
        valido, msg = self.validar_arquivo(caminho)
        if not valido:
            return {"status": "ERROR", "error": msg}
        
        if not Document:
            return {"status": "ERROR", "error": "Biblioteca python-docx não instalada."}

        try:
            doc = Document(caminho)
            texto = "\n".join([para.text for para in doc.paragraphs])
            return {
                "status": "SUCCESS",
                "conteudo": texto,
                "metadados": {"paragrafos": len(doc.paragraphs), "formato": "DOCX"}
            }
        except Exception as e:
            return {"status": "ERROR", "error": str(e)}

def executar(entrada):
    return ExtratorDOCX().executar(entrada)
